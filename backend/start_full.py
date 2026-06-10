"""完整系统启动脚本 - 修复版，支持文档解析"""
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
from typing import Optional, List, Dict
import uvicorn
import sys
import os
import json
import requests
from datetime import datetime
import io
import re

# 添加当前目录到 Python 路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 创建完整的 FastAPI 应用
app = FastAPI(
    title="AI 知识库管理系统",
    description="个人知识库管理系统",
    version="2.2.0"
)

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 数据存储
documents_db = []
chat_history = []
knowledge_units = []  # 知识单元列表（核心知识块）
document_chunks = {}  # 文档内容块（文档ID -> 知识片段列表）


# AI配置
AI_PROVIDER = "siliconflow"
AI_BASE_URL = "https://api.siliconflow.cn/v1"
AI_MODEL = "Qwen/Qwen2.5-7B-Instruct"
AI_API_KEY = ""

# 向量数据库配置（简化版，使用内存存储）
document_vectors = {}

# 尝试加载 .env 文件
try:
    env_file = os.path.join(os.path.dirname(__file__), '.env')
    if os.path.exists(env_file):
        with open(env_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    key = key.strip()
                    value = value.strip().strip('"\'')
                    os.environ[key] = value
        print("✅ 配置文件加载成功")
    else:
        print("⚠️  未找到 .env 文件")
except Exception as e:
    print(f"⚠️  加载配置文件失败: {e}")

# 从环境变量读取配置
AI_API_KEY = os.environ.get('OPENAI_API_KEY', '')
if not AI_API_KEY:
    print("⚠️  未配置 API Key，部分功能将受限")

# 数据模型
class SearchQuery(BaseModel):
    query: str

class GenerateContentQuery(BaseModel):
    topic: str
    style: str = "professional"

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: List[ChatMessage]

# ============ 文档解析功能 ============

def parse_pdf(content: bytes) -> str:
    """解析 PDF 文件"""
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text.strip()
    except:
        # 如果 PyMuPDF 不可用，尝试基础方法
        try:
            text = content.decode('utf-8', errors='ignore')
            text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
            return text.strip()[:10000]
        except:
            return ""

def parse_word(content: bytes) -> str:
    """解析 Word 文档"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    text += " | ".join(row_text) + "\n"
        return text.strip()
    except:
        return ""

def parse_markdown(content: bytes) -> str:
    """解析 Markdown 文件"""
    try:
        return content.decode('utf-8', errors='ignore').strip()
    except:
        return ""

def parse_text(content: bytes) -> str:
    """解析文本文件"""
    for encoding in ['utf-8', 'gbk', 'gb2312']:
        try:
            return content.decode(encoding, errors='ignore').strip()
        except:
            continue
    return ""

def parse_document(content: bytes, filename: str) -> str:
    """统一的文档解析接口"""
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    
    # 解析器映射
    parsers = {
        'pdf': parse_pdf,
        'doc': parse_word,
        'docx': parse_word,
        'md': parse_markdown,
        'txt': parse_text,
        'markdown': parse_markdown,
    }
    
    parser = parsers.get(ext)
    if parser:
        try:
            result = parser(content)
            print(f"✅ 成功解析 {filename} ({ext}), 长度: {len(result)}")
            return result
        except Exception as e:
            print(f"❌ 解析 {filename} 失败: {e}")
            return ""
    
    # 默认尝试文本解析
    return parse_text(content)

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """将文本分块"""
    if len(text) <= chunk_size:
        return [text] if text.strip() else []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end < len(text):
            for sep in ['\n\n', '\n', '。', '！', '？', '. ', '! ', '? ']:
                sep_pos = text.rfind(sep, start, end)
                if sep_pos > start:
                    end = sep_pos + len(sep)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        if start <= (chunks[-1].find(chunk[-50:]) if chunks else 0):
            start = end
    
    return [c for c in chunks if c]

# ============ API 端点 ============

@app.get("/", response_class=HTMLResponse)
async def read_root():
    """返回前端主页"""
    index_path = os.path.join(os.path.dirname(__file__), "index.html")
    if os.path.exists(index_path):
        with open(index_path, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>AI 知识库系统</h1><p>前端页面加载中...</p>")

@app.get("/api/health")
async def health_check():
    return {
        "status": "healthy",
        "service": "ai-knowledge-base",
        "version": "2.2.0",
        "api_configured": bool(AI_API_KEY),
        "documents_count": len(documents_db),
        "features": ["文档解析", "内容索引", "语义搜索", "AI对话"]
    }

@app.get("/api/status")
async def get_status():
    return {
        "status": "running",
        "version": "2.2.0",
        "ai_provider": AI_PROVIDER,
        "ai_model": AI_MODEL,
        "api_configured": bool(AI_API_KEY),
        "documents_count": len(documents_db),
        "indexed_chunks": sum(len(v.get('chunks', [])) for v in document_vectors.values())
    }

# 文档管理
@app.get("/api/documents")
async def list_documents():
    return {
        "total": len(documents_db),
        "documents": documents_db
    }

@app.post("/api/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    try:
        # 检查文件类型
        allowed_extensions = ['.pdf', '.doc', '.docx', '.txt', '.md', '.markdown']
        filename_lower = file.filename.lower() if file.filename else ''
        
        is_allowed = any(filename_lower.endswith(ext) for ext in allowed_extensions)
        if not is_allowed:
            return {
                "success": False,
                "message": f"不支持的文件类型，请上传 PDF、Word、TXT 或 Markdown 文件"
            }
        
        content = await file.read()
        file_size = len(content)
        
        # 解析文档内容
        text_content = parse_document(content, file.filename)
        
        # 创建文档记录（不立即进行知识提取）
        doc = {
            "id": len(documents_db) + 1,
            "name": file.filename,
            "size": file_size,
            "uploaded_at": datetime.now().isoformat(),
            "content_length": len(text_content),
            "file_type": get_file_type(file.filename),
            "knowledge_processed": False,  # 标记是否已处理
            "knowledge_unit_ids": []  # 关联的知识单元ID列表
        }
        documents_db.append(doc)
        
        # 索引文档内容（只做基本分词）
        if text_content and len(text_content) > 10:
            chunks = chunk_text(text_content)
            document_vectors[doc["id"]] = {
                "content": text_content,
                "chunks": chunks,
                "filename": file.filename
            }
            message = f"文档 {file.filename} 上传成功，已解析 {len(chunks)} 个内容块。请点击「更新知识库」提取知识。"
        else:
            message = f"文档 {file.filename} 上传成功（内容解析失败或为空）"
        
        return {
            "success": True,
            "document": doc,
            "message": message,
            "parsed": bool(text_content and len(text_content) > 10),
            "chunk_count": len(document_vectors.get(doc["id"], {}).get('chunks', []))
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"上传失败: {str(e)}"
        }

# 后台任务状态存储
tasks_db = {}

import threading

def run_knowledge_update_in_background(task_id: str):
    """后台线程：执行知识库更新"""
    global tasks_db
    
    tasks_db[task_id] = {
        "status": "running",
        "start_time": datetime.now().isoformat(),
        "message": "正在更新知识库...",
        "progress": 0,
        "total": 0,
        "processed": 0,
        "units_created": 0
    }
    
    try:
        if not AI_API_KEY:
            tasks_db[task_id]["status"] = "error"
            tasks_db[task_id]["message"] = "未配置 API Key，无法更新知识库"
            return
        
        unprocessed_docs = [doc for doc in documents_db if not doc.get("knowledge_processed", False)]
        
        if not unprocessed_docs:
            tasks_db[task_id]["status"] = "completed"
            tasks_db[task_id]["message"] = "所有文档都已处理完毕"
            tasks_db[task_id]["progress"] = 100
            return
        
        total_docs = len(unprocessed_docs)
        tasks_db[task_id]["total"] = total_docs
        
        total_units_created = 0
        processed_count = 0
        
        for i, doc in enumerate(unprocessed_docs):
            doc_id = doc["id"]
            vector_data = document_vectors.get(doc_id)
            
            # 更新进度
            tasks_db[task_id]["progress"] = int((i / total_docs) * 100)
            tasks_db[task_id]["message"] = f"正在处理第 {i + 1}/{total_docs} 个文档..."
            
            if not vector_data:
                doc["knowledge_processed"] = True
                processed_count += 1
                tasks_db[task_id]["processed"] = processed_count
                continue
            
            content = vector_data.get("content", "")
            filename = vector_data.get("filename", "")
            
            if not content or len(content) < 100:
                doc["knowledge_processed"] = True
                processed_count += 1
                tasks_db[task_id]["processed"] = processed_count
                continue
            
            # 调用智能拆分函数（直接调用同步版本）
            try:
                assigned_units = smart_split_document_sync(doc_id, content, filename)
            except Exception as inner_e:
                print(f"处理文档 {doc_id} 时出错: {inner_e}")
                doc["knowledge_processed"] = True
                processed_count += 1
                tasks_db[task_id]["processed"] = processed_count
                continue
            
            doc["knowledge_processed"] = True
            doc["knowledge_unit_ids"] = assigned_units
            
            total_units_created += len(assigned_units)
            processed_count += 1
            tasks_db[task_id]["processed"] = processed_count
            tasks_db[task_id]["units_created"] = total_units_created
        
        tasks_db[task_id]["status"] = "completed"
        tasks_db[task_id]["progress"] = 100
        tasks_db[task_id]["message"] = f"知识库更新完成！处理了 {processed_count} 个文档，创建/更新了 {total_units_created} 个知识单元"
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        tasks_db[task_id]["status"] = "error"
        tasks_db[task_id]["message"] = f"更新失败: {str(e)}"

@app.post("/api/knowledge/update")
async def update_knowledge_base():
    """启动知识库更新（后台运行）"""
    import uuid
    task_id = str(uuid.uuid4())
    
    # 检查是否有正在运行的任务
    for tid, task in tasks_db.items():
        if task["status"] == "running":
            return {
                "success": False,
                "message": "已有更新任务在运行中，请等待完成",
                "task_id": tid,
                "status": task["status"]
            }
    
    # 启动后台线程
    thread = threading.Thread(target=run_knowledge_update_in_background, args=(task_id,), daemon=True)
    thread.start()
    
    return {
        "success": True,
        "message": "知识库更新任务已启动，请在后台查看进度",
        "task_id": task_id,
        "status": "running"
    }

@app.get("/api/knowledge/task/{task_id}")
async def get_task_status(task_id: str):
    """获取后台任务状态"""
    task = tasks_db.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    return {
        "success": True,
        "task_id": task_id,
        "status": task.get("status", "unknown"),
        "message": task.get("message", ""),
        "progress": task.get("progress", 0),
        "total": task.get("total", 0),
        "processed": task.get("processed", 0),
        "units_created": task.get("units_created", 0),
        "start_time": task.get("start_time", "")
    }

@app.get("/api/knowledge/status")
async def get_knowledge_status():
    """获取知识库状态"""
    total_docs = len(documents_db)
    processed_docs = sum(1 for doc in documents_db if doc.get("knowledge_processed", False))
    unprocessed_docs = total_docs - processed_docs
    
    return {
        "success": True,
        "total_documents": total_docs,
        "processed_documents": processed_docs,
        "unprocessed_documents": unprocessed_docs,
        "total_knowledge_units": len(knowledge_units)
    }


# 常见的停用词和垃圾词
STOPWORDS = {
    'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
    'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
    'this', 'that', 'these', 'those', 'it', 'its', 'as', 'if', 'then',
    'docx', 'pdf', 'txt', 'doc', 'char', '字符', '文档', '文件', '片段',
    '概述', '简介', '说明', '内容', '部分', '章节', '等等', '一些', '的话',
    '关于', '对于', '以及', '可以', '能够', '需要', '应该', '可能', '例如',
    '包括', '涉及', '通过', '使用', '方法', '技术', '系统', '功能', '模块',
    '我们', '你们', '他们', '什么', '怎么', '为什么', '如何', '哪里', '哪个'
}

def extract_keywords(text: str, max_keywords: int = 5) -> list:
    """从文本中提取关键词（更干净的版本）"""
    if not text:
        return []
    import re
    from collections import Counter

    # 去除所有标点符号，仅保留中文、英文、数字
    clean_text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', text)

    # 提取2-4字中文词
    chinese_words = []
    for length in [4, 3, 2]:
        for i in range(len(clean_text) - length + 1):
            word = clean_text[i:i+length]
            if re.match(r'^[\u4e00-\u9fa5]{2,4}$', word) and word not in STOPWORDS:
                chinese_words.append(word)

    # 提取英文单词（3-15字符）和数字（4+位）
    english_words = re.findall(r'[a-zA-Z]{3,15}', text)
    numbers = re.findall(r'\d{4,}', text)

    all_words = []
    all_words.extend(chinese_words)
    for w in english_words:
        if w.lower() not in STOPWORDS:
            all_words.append(w)
    all_words.extend(numbers[:2])

    # 统计词频
    word_counts = Counter(all_words)
    sorted_words = sorted(word_counts.items(), key=lambda x: (x[1], len(x[0])), reverse=True)

    # 选取高频词并去除子词重复
    keywords = []
    for word, count in sorted_words:
        if len(keywords) >= max_keywords:
            break
        # 检查是否已有更长的词包含当前词，或当前词包含某个已有关键词
        is_redundant = False
        for kw in keywords:
            if word in kw or kw in word:
                is_redundant = True
                break
        if not is_redundant:
            keywords.append(word)

    return keywords[:max_keywords]

def extract_key_sentences(text: str, max_sentences: int = 3, max_chars_per_sentence: int = 400) -> list:
    """从文本中提取完整的关键句子，避免内容被截断"""
    if not text:
        return []
    import re

    # 按中英文句号、问号、感叹号、分号等切割句子，保留标点
    raw_sentences = re.split(r'(?<=[。！？；!?;])\s*|[\n\r]+', text)
    sentences = []
    for s in raw_sentences:
        s = s.strip()
        if 15 <= len(s) <= max_chars_per_sentence:
            sentences.append(s)

    # 如果句子过少，尝试按逗号/顿号补充
    if len(sentences) < max_sentences:
        extra_parts = re.split(r'[，、,，]', text)
        for p in extra_parts:
            p = p.strip()
            if 20 <= len(p) <= max_chars_per_sentence and p not in sentences:
                sentences.append(p)

    # 打分：优先选择含有关键词、含数字/特定术语的句子
    keywords = extract_keywords(text, max_keywords=5)
    keyword_set = set(keywords)

    scored = []
    for s in sentences:
        score = 0
        for kw in keyword_set:
            if kw in s:
                score += 1
        # 数字加分
        score += len(re.findall(r'\d+', s)) * 0.3
        # 长度适中加分（100-300字）
        if 50 <= len(s) <= 300:
            score += 1
        scored.append((score, s))

    scored.sort(key=lambda x: x[0], reverse=True)
    result = [s for _, s in scored[:max_sentences]]

    # 如果还是没有足够内容，取整段开头（保持完整，不强行截断）
    if not result and len(text) > 20:
        # 寻找第一个句号位置，最多取到该处
        first_end = re.search(r'[。！？!?;；]', text)
        if first_end:
            candidate = text[:first_end.end()].strip()
            if len(candidate) <= max_chars_per_sentence:
                result.append(candidate)
            else:
                result.append(text[:max_chars_per_sentence].strip())
        else:
            result.append(text[:max_chars_per_sentence].strip())

    return result

def calculate_content_similarity(content1: str, content2: str) -> float:
    """计算两个内容片段的相似度（基于字符n-gram重叠）"""
    if not content1 or not content2:
        return 0.0
    import re
    
    def get_ngrams(text, n=3):
        ngrams = set()
        text = re.sub(r'[\s\n\r\t]+', '', text)
        for i in range(len(text) - n + 1):
            ngrams.add(text[i:i+n])
        return ngrams
    
    ngrams1 = get_ngrams(content1[:500])
    ngrams2 = get_ngrams(content2[:500])
    
    if not ngrams1 or not ngrams2:
        return 0.0
    
    intersection = len(ngrams1 & ngrams2)
    union = len(ngrams1 | ngrams2)
    
    if union == 0:
        return 0.0
    return intersection / union

def generate_title_from_content(content: str, filename: str) -> str:
    """基于内容生成简洁标题：选1-3个最核心关键词组合"""
    keywords = extract_keywords(content, max_keywords=5)
    if keywords:
        # 选取1-3个最核心关键词，用空格或自然分隔
        num = min(3, len(keywords))
        if num == 1:
            return keywords[0]
        return ' '.join(keywords[:num])
    clean_name = filename.replace('.docx', '').replace('.pdf', '').replace('.txt', '')
    return clean_name[:20]

def smart_split_document_sync(doc_id: int, content: str, filename: str) -> list:
    """同步版本：智能拆分文档内容到知识单元"""
    print(f"\n=== 开始处理文档: {filename} ===")
    print(f"文档ID: {doc_id}, 内容长度: {len(content)}")
    
    if not content or len(content) < 100:
        print(f"跳过：内容太短")
        return []
    
    try:
        import re
        # 将内容分块
        text_chunks = chunk_text(content, chunk_size=500, overlap=50)
        print(f"生成了 {len(text_chunks)} 个文本块")
        
        # 为每个文本块提取内容和关键词
        chunk_units = []
        
        # 处理前5个块
        for i, chunk in enumerate(text_chunks[:5]):
            if len(chunk) < 100:
                continue

            # 从内容中提取关键词（不依赖文件名）
            content_keywords = extract_keywords(chunk, max_keywords=5)

            # 提取完整的关键句子作为内容（避免截断）
            key_sentences = extract_key_sentences(chunk, max_sentences=3)
            sentence_content = ' '.join(key_sentences) if key_sentences else chunk[:400]

            # 生成基于内容的摘要
            summary = generate_title_from_content(chunk, filename)

            chunk_units.append({
                "content": sentence_content,
                "summary": summary,
                "keywords": content_keywords
            })
            print(f"  块 {i+1}: 提取关键词 {content_keywords}")

        # 如果提取失败，用后备方案
        if not chunk_units:
            print("  提取失败，使用后备方案")
            for i, chunk in enumerate(text_chunks[:3]):
                if len(chunk) > 100:
                    keywords = extract_keywords(chunk, max_keywords=5)
                    key_sentences = extract_key_sentences(chunk, max_sentences=3)
                    sentence_content = ' '.join(key_sentences) if key_sentences else chunk[:400]
                    chunk_units.append({
                        "content": sentence_content,
                        "summary": generate_title_from_content(chunk, filename),
                        "keywords": keywords
                    })
        
        print(f"  准备分配 {len(chunk_units)} 个内容片段到知识单元")
        
        # 分配到知识单元（基于内容相似度判断）
        assigned_unit_ids = []
        created_count = 0
        
        for idx, unit_data in enumerate(chunk_units[:3]):
            unit_content = unit_data["content"]
            unit_summary = unit_data["summary"]
            unit_keywords = unit_data["keywords"]
            
            if not unit_content or not unit_keywords:
                continue
            
            print(f"  处理片段 {idx+1}: 摘要='{unit_summary}', 关键词={unit_keywords}")
            
            # 查找相似的现有知识单元
            matched_unit = None
            best_similarity = 0.0
            
            for existing_unit in knowledge_units:
                existing_keywords = set(existing_unit.get("keywords", []))
                current_keywords = set(unit_keywords)
                keyword_overlap = len(existing_keywords & current_keywords)
                
                content_sim = 0.0
                for existing_content in existing_unit.get("contents", []):
                    sim = calculate_content_similarity(unit_content, existing_content.get("content", ""))
                    content_sim = max(content_sim, sim)
                
                combined_score = 0.0
                if existing_keywords and current_keywords:
                    keyword_score = keyword_overlap / min(len(existing_keywords), len(current_keywords))
                    combined_score = keyword_score * 0.5 + content_sim * 0.5
                
                print(f"    和单元 #{existing_unit['id']} 相似度: {combined_score:.3f}")
                
                if combined_score > best_similarity and combined_score > 0.15:
                    best_similarity = combined_score
                    matched_unit = existing_unit
            
            if matched_unit:
                print(f"  -> 加入现有单元 #{matched_unit['id']} (相似度 {best_similarity:.3f})")
                matched_unit["contents"].append({
                    "content": unit_content,
                    "source_doc_id": doc_id,
                    "source_filename": filename
                })
                merged_keywords = list(set(matched_unit.get("keywords", []) + unit_keywords))[:10]
                matched_unit["keywords"] = merged_keywords
                matched_unit["last_updated"] = datetime.now().isoformat()
                assigned_unit_ids.append(matched_unit["id"])
            else:
                new_unit = {
                    "id": len(knowledge_units) + 1,
                    "summary": unit_summary,
                    "keywords": unit_keywords,
                    "contents": [{
                        "content": unit_content,
                        "source_doc_id": doc_id,
                        "source_filename": filename
                    }],
                    "created_at": datetime.now().isoformat(),
                    "last_updated": datetime.now().isoformat()
                }
                knowledge_units.append(new_unit)
                assigned_unit_ids.append(new_unit["id"])
                created_count += 1
                print(f"  -> 创建新单元 #{new_unit['id']}: {unit_summary}")
        
        print(f"=== 完成处理: 创建 {created_count} 个新单元, 加入现有 {len(assigned_unit_ids) - created_count} 个单元 ===\n")
        return list(set(assigned_unit_ids))
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"处理文档失败: {e}")
        return []

# 知识单元管理接口
@app.get("/api/knowledge/units")
async def get_knowledge_units():
    """获取所有知识单元"""
    return {
        "success": True,
        "units": knowledge_units,
        "total": len(knowledge_units)
    }

@app.get("/api/knowledge/units/{unit_id}")
async def get_unit_detail(unit_id: int):
    """获取知识单元详情"""
    unit = next((u for u in knowledge_units if u["id"] == unit_id), None)
    if not unit:
        raise HTTPException(status_code=404, detail="知识单元不存在")
    
    # 获取来源文档信息
    source_docs = []
    for content in unit.get("contents", []):
        doc_id = content.get("source_doc_id")
        doc = next((d for d in documents_db if d["id"] == doc_id), None)
        if doc and doc not in source_docs:
            source_docs.append({
                "id": doc["id"],
                "name": doc.get("name", "")
            })
    
    return {
        "success": True,
        "unit": unit,
        "source_documents": source_docs
    }

@app.delete("/api/knowledge/units/{unit_id}")
async def delete_unit(unit_id: int):
    """删除知识单元"""
    global knowledge_units
    
    unit = next((u for u in knowledge_units if u["id"] == unit_id), None)
    if not unit:
        raise HTTPException(status_code=404, detail="知识单元不存在")
    
    knowledge_units = [u for u in knowledge_units if u["id"] != unit_id]
    
    # 从文档中移除关联
    for doc in documents_db:
        if unit_id in doc.get("knowledge_unit_ids", []):
            doc["knowledge_unit_ids"].remove(unit_id)
    
    return {"success": True, "message": "知识单元已删除"}

@app.post("/api/knowledge/units/{unit_id}/refresh")
async def refresh_unit(unit_id: int):
    """刷新知识单元的摘要和关键词"""
    global knowledge_units
    
    unit = next((u for u in knowledge_units if u["id"] == unit_id), None)
    if not unit:
        raise HTTPException(status_code=404, detail="知识单元不存在")
    
    if not AI_API_KEY or not unit.get("contents"):
        return {"success": True, "message": "无法刷新"}
    
    try:
        # 合并所有内容
        all_content = "\n\n".join([c.get("content", "") for c in unit.get("contents", [])])
        
        system_prompt = """你是一个知识整理专家。请分析以下内容，生成一个简洁的知识单元摘要。

请以JSON格式返回：
{
    "summary": "用一句话概括这个知识点（15-30字）",
    "keywords": ["关键词1", "关键词2", "关键词3", "关键词4"]
}

只返回JSON。"""
        
        ai_response = call_ai_api([
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": all_content[:2000]}
        ])
        
        try:
            result_json = ai_response[ai_response.find('{') : ai_response.rfind('}') + 1]
            result = json.loads(result_json)
            unit["summary"] = result.get("summary", unit.get("summary"))
            unit["keywords"] = result.get("keywords", unit.get("keywords", []))
            unit["last_updated"] = datetime.now().isoformat()
        except:
            pass
        
        return {"success": True, "message": "知识单元已刷新"}
    except Exception as e:
        return {"success": False, "message": f"刷新失败: {str(e)}"}

@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: int):
    global documents_db
    original_count = len(documents_db)
    documents_db = [d for d in documents_db if d.get("id") != doc_id]
    
    if doc_id in document_vectors:
        del document_vectors[doc_id]
    
    if len(documents_db) < original_count:
        return {"success": True, "message": "文档已删除"}
    else:
        raise HTTPException(status_code=404, detail="文档不存在")

@app.get("/api/documents/{doc_id}/content")
async def get_document_content(doc_id: int):
    """获取文档内容"""
    if doc_id in document_vectors:
        vector_data = document_vectors[doc_id]
        return {
            "success": True,
            "filename": vector_data["filename"],
            "content": vector_data["content"][:5000] if vector_data.get("content") else "",
            "chunks": vector_data["chunks"][:10] if vector_data.get("chunks") else []
        }
    
    return {
        "success": True,
        "filename": "",
        "content": "",
        "chunks": []
    }

def get_file_type(filename: str) -> str:
    """获取文件类型"""
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    types = {
        'pdf': 'PDF 文档',
        'doc': 'Word 文档 (2003)',
        'docx': 'Word 文档',
        'ppt': 'PPT 演示文稿 (2003)',
        'pptx': 'PPT 演示文稿',
        'xls': 'Excel 表格 (2003)',
        'xlsx': 'Excel 表格',
        'txt': '文本文档',
        'md': 'Markdown 文档',
        'html': 'HTML 网页',
        'json': 'JSON 数据文件',
        'xml': 'XML 数据文件',
        'csv': 'CSV 数据文件'
    }
    return types.get(ext, '未知类型')

# 搜索功能 - 只返回文档搜索结果，不调用AI
@app.post("/api/search")
async def search_knowledge(query: SearchQuery):
    try:
        search_results = []
        query_lower = query.query.lower()
        
        # 匹配文件名
        for doc in documents_db:
            doc_name = doc.get('name', '').lower()
            doc_type = doc.get('file_type', '').lower()
            
            if query_lower in doc_name or query_lower in doc_type:
                search_results.append({
                    "type": "filename",
                    "doc_id": doc["id"],
                    "filename": doc.get("name", ""),
                    "file_type": doc.get("file_type", ""),
                    "file_size": doc.get("size", 0),
                    "uploaded_at": doc.get("uploaded_at", ""),
                    "relevance": 0.9,
                    "content": f"文件名匹配：{doc.get('name', '')}",
                    "chunks": []
                })
        
        # 搜索文档内容
        for doc_id, vector_data in document_vectors.items():
            filename = vector_data.get("filename", "")
            chunks = vector_data.get("chunks", [])
            full_content = vector_data.get("content", "")
            
            if query_lower in full_content.lower():
                matching_chunks = []
                for i, chunk in enumerate(chunks):
                    if query_lower in chunk.lower():
                        matching_chunks.append({
                            "chunk_index": i,
                            "content": chunk
                        })
                
                if matching_chunks:
                    # 查找文档信息
                    doc_info = next((d for d in documents_db if d["id"] == doc_id), None)
                    search_results.append({
                        "type": "content",
                        "doc_id": doc_id,
                        "filename": filename,
                        "file_type": doc_info.get("file_type", "") if doc_info else "",
                        "file_size": doc_info.get("size", 0) if doc_info else 0,
                        "uploaded_at": doc_info.get("uploaded_at", "") if doc_info else "",
                        "relevance": 0.8,
                        "content": full_content[:200] + "..." if len(full_content) > 200 else full_content,
                        "chunks": matching_chunks
                    })
        
        # 按相关度排序
        search_results.sort(key=lambda x: x["relevance"], reverse=True)
        
        return {
            "success": True,
            "query": query.query,
            "total_matches": len(search_results),
            "results": search_results
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"搜索失败: {str(e)}",
            "total_matches": 0,
            "results": []
        }

# 对话式 AI 助手（集成内容创作和决策辅助）
@app.post("/api/chat")
async def chat_with_assistant(chat_request: ChatRequest):
    try:
        if not AI_API_KEY:
            return {
                "success": False,
                "message": "未配置 API Key，请先配置",
                "response": ""
            }

        # 获取最后一条用户消息，判断任务类型
        user_message = chat_request.messages[-1].content if chat_request.messages else ""
        
        # 判断是否是创作或决策辅助请求
        is_creation_request = any(keyword in user_message for keyword in ["创作", "写一篇", "生成内容", "帮我写", "内容生成"])
        is_decision_request = any(keyword in user_message for keyword in ["决策", "分析", "建议", "对比", "选择", "方案"])
        
        # 构建系统提示
        system_prompt = """你是一个智能知识库助手，可以进行多轮对话。你的知识库中包含了用户上传的文档。"""
        
        if documents_db:
            system_prompt += f"\n\n目前知识库中有 {len(documents_db)} 个文档：\n"
            for i, doc in enumerate(documents_db, 1):
                has_content = doc["id"] in document_vectors
                content_info = f"（已解析 {len(document_vectors.get(doc['id'], {}).get('chunks', []))} 块）" if has_content else "（未解析）"
                system_prompt += f"{i}. 📄 {doc.get('name', '')} {content_info}\n"
            
            if document_vectors:
                system_prompt += "\n📖 文档内容摘要：\n"
                for doc_id, vector_data in list(document_vectors.items())[:3]:
                    content = vector_data.get("content", "")[:500]
                    if content:
                        system_prompt += f"\n【{vector_data.get('filename', '未知文件')}】\n"
                        system_prompt += f"{content}...\n"
        else:
            system_prompt += "\n\n目前知识库是空的。"
        
        # 根据任务类型调整提示
        if is_creation_request:
            system_prompt += """
你的职责：
1. 基于知识库内容进行内容创作
2. 结合相关文档生成高质量内容
3. 注明引用来源
4. 使用清晰的结构组织内容
5. 支持多种风格（专业、轻松、学术等）

请用中文创作，内容要有深度和价值。"""
        elif is_decision_request:
            system_prompt += """
你的职责：
1. 梳理问题背景和分析需求
2. 列出可行的方案选项
3. 分析各方案的优缺点
4. 给出推荐建议
5. 提供风险提示
6. 基于知识库内容给出有依据的分析

请用结构化的方式输出，帮助用户做出明智决策。"""
        else:
            system_prompt += """
你的职责：
1. 友好地与用户对话
2. 基于知识库中的文档内容回答用户的问题
3. 如果知识库中有相关内容，尽量引用
4. 如果没有相关文档，诚实告知并建议上传
5. 可以进行多轮对话，理解上下文

请用中文回答，保持专业和友好。"""

        messages = [{"role": "system", "content": system_prompt}]
        messages.extend([{"role": m.role, "content": m.content} for m in chat_request.messages])
        
        ai_response = call_ai_api(messages)
        
        return {
            "success": True,
            "response": ai_response
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"对话失败: {str(e)}",
            "response": ""
        }

# AI API 调用
def call_ai_api(messages: List[Dict]) -> str:
    """调用硅基流动 API"""
    if not AI_API_KEY:
        return "未配置 API Key，无法进行 AI 交互。"
    
    url = f"{AI_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {AI_API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": AI_MODEL,
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(url, headers=headers, json=data, timeout=60)
        response.raise_for_status()
        
        result = response.json()
        
        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"]
        else:
            return f"API 返回格式异常: {str(result)}"
            
    except requests.exceptions.Timeout:
        return "请求超时，请稍后重试。"
    except requests.exceptions.RequestException as e:
        return f"API调用失败: {str(e)}\n\n请检查 API Key 是否正确配置。"
    except Exception as e:
        return f"处理响应失败: {str(e)}"

if __name__ == "__main__":
    print("=" * 70)
    print("🚀  AI 知识库管理系统 v2.2 - 修复版启动")
    print("=" * 70)
    print("\n")
    print("🌐  前端访问:  http://localhost:8000")
    print("📚  API 文档:  http://localhost:8000/docs")
    print("🔍  健康检查: http://localhost:8000/api/health")
    print("🤖  AI 提供商: 硅基流动 (SiliconFlow)")
    print("📄  支持格式: PDF, Word, Markdown, TXT 等")
    print("💬  功能: 文档解析 | 语义搜索 | AI 对话")
    print("=" * 70)
    print("\n")
    
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        reload=False,
        log_level="info"
    )
